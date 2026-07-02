from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


TABLE_INPUTS = {
    "../outputs/paper_assets/main_calibration_5scenario.tex": ROOT
    / "outputs/paper_assets/main_calibration_5scenario.tex",
    "../outputs/paper_assets/main_calibration_progression_5scenario.tex": ROOT
    / "outputs/paper_assets/main_calibration_progression_5scenario.tex",
    "../outputs/paper_assets/robustness_calibration_10scenario.tex": ROOT
    / "outputs/paper_assets/robustness_calibration_10scenario.tex",
    "../outputs/paper_assets/rotating_cv_calibration_5scenario.tex": ROOT
    / "outputs/paper_assets/rotating_cv_calibration_5scenario.tex",
    "../outputs/paper_assets/matched_non_calibrated_5scenario.tex": ROOT
    / "outputs/paper_assets/matched_non_calibrated_5scenario.tex",
    "../outputs/paper_assets/short_new_f1_table.tex": ROOT
    / "outputs/paper_assets/short_new_f1_table.tex",
    "../outputs/paper_assets_supervised_metric/main_calibration_5scenario.tex": ROOT
    / "outputs/paper_assets_supervised_metric/main_calibration_5scenario.tex",
    "../outputs/paper_assets_supervised_metric/robustness_calibration_10scenario.tex": ROOT
    / "outputs/paper_assets_supervised_metric/robustness_calibration_10scenario.tex",
    "../outputs/paper_assets_supervised_metric/rotating_cv_calibration_5scenario.tex": ROOT
    / "outputs/paper_assets_supervised_metric/rotating_cv_calibration_5scenario.tex",
    "../outputs/paper_assets_supervised_metric/matched_non_calibrated_5scenario.tex": ROOT
    / "outputs/paper_assets_supervised_metric/matched_non_calibrated_5scenario.tex",
    "../outputs/paper_assets_supervised_metric/llm_contribution_ablation.tex": ROOT
    / "outputs/paper_assets_supervised_metric/llm_contribution_ablation.tex",
    "../outputs/paper_assets_supervised_metric/bootstrap_ci_summary.tex": ROOT
    / "outputs/paper_assets_supervised_metric/bootstrap_ci_summary.tex",
    "../outputs/paper_assets_supervised_metric/latency_accounting.tex": ROOT
    / "outputs/paper_assets_supervised_metric/latency_accounting.tex",
}

FIGURE_INPUTS = {
    "../outputs/paper_assets/system_architecture.tex": {
        "image": ROOT / "outputs/paper_assets/system_architecture.png",
        "caption": (
            "Physics-warm-started multimodal LLM calibration architecture. "
            "The LLM supplies structured correction evidence, while deterministic "
            "tools, guardrails, and a learned pixel-level calibrator produce the "
            "final operational burn map."
        ),
    }
}


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def strip_tex(text: str) -> str:
    text = text.replace("~", " ")
    text = text.replace("\\%", "%")
    text = text.replace("\\_", "_")
    text = text.replace("$", "")
    text = text.replace("``", '"').replace("''", '"')
    text = re.sub(r"\\texttt\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\(?:auto)?ref\{([^{}]*)\}", "??", text)
    text = re.sub(r"\\label\{[^{}]*\}", "", text)
    text = re.sub(r"\\url\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def replace_cross_references(text: str, label_refs: dict[str, str]) -> str:
    def repl(match: re.Match[str]) -> str:
        return label_refs.get(match.group(1), "??")

    return re.sub(r"\\(?:auto)?ref\{([^{}]+)\}", repl, text)


def labels_in_text(text: str) -> list[str]:
    return re.findall(r"\\label\{([^{}]+)\}", text)


def collect_label_refs(tex: str, tex_path: Path) -> dict[str, str]:
    label_refs: dict[str, str] = {}
    figure_index = 1
    table_index = 1
    in_figure = False
    figure_block: list[str] = []

    for line in tex.splitlines():
        stripped = line.strip()
        if stripped.startswith("\\begin{figure}"):
            in_figure = True
            figure_block = [stripped]
            continue
        if in_figure:
            figure_block.append(stripped)
            if stripped.startswith("\\end{figure}"):
                for label in labels_in_text("\n".join(figure_block)):
                    label_refs[label] = str(figure_index)
                figure_index += 1
                in_figure = False
            continue

        input_match = re.match(r"\\input\{([^{}]+)\}", stripped)
        if not input_match:
            continue
        input_key = input_match.group(1)
        if input_key in FIGURE_INPUTS:
            input_path = (tex_path.parent / input_key).resolve()
            if input_path.exists():
                for label in labels_in_text(read_text(input_path)):
                    label_refs[label] = str(figure_index)
            figure_index += 1
        elif input_key in TABLE_INPUTS:
            table_path = TABLE_INPUTS[input_key]
            if table_path.exists():
                for label in labels_in_text(read_text(table_path)):
                    label_refs[label] = str(table_index)
            table_index += 1

    return label_refs


def parse_bib_order(bbl_path: Path) -> tuple[dict[str, int], list[str]]:
    if not bbl_path.exists():
        return {}, []
    bbl = read_text(bbl_path)
    keys = re.findall(r"\\bibitem(?:\[[^\]]*\])?\{([^{}]+)\}", bbl)
    key_to_num = {key: idx + 1 for idx, key in enumerate(keys)}
    chunks = re.split(r"\\bibitem(?:\[[^\]]*\])?\{[^{}]+\}", bbl)
    entries: list[str] = []
    for chunk in chunks[1:]:
        cleaned = strip_tex(chunk)
        cleaned = cleaned.replace("newblock", "")
        cleaned = cleaned.replace("thebibliography", "")
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        if cleaned:
            entries.append(cleaned)
    return key_to_num, entries


def replace_citations(text: str, key_to_num: dict[str, int]) -> str:
    def repl(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",")]
        nums = [str(key_to_num[key]) for key in keys if key in key_to_num]
        return "[" + ",".join(nums) + "]" if nums else ""

    return re.sub(r"\\citep\{([^{}]+)\}", repl, text)


def parse_table(table_path: Path) -> tuple[list[list[str]], str]:
    text = read_text(table_path)
    caption_match = re.search(r"\\caption\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}", text, re.S)
    caption = strip_tex(caption_match.group(1)) if caption_match else ""
    rows: list[list[str]] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line.startswith("\\") or "&" not in line:
            continue
        line = line.replace("\\\\", "").strip()
        cells = [strip_tex(cell).replace("_", " ") for cell in line.split("&")]
        rows.append(cells)
    if len(rows) > 1 and rows[0][0] == "Run":
        run_values = {row[0] for row in rows[1:] if row}
        if len(run_values) == 1:
            rows = [row[1:] for row in rows]
    return rows, caption


def add_table(doc: Document, table_path: Path, table_index: int) -> None:
    rows, caption = parse_table(table_path)
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for col_idx, value in enumerate(rows[0]):
        cell = table.rows[0].cells[col_idx]
        cell.text = value
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7)
    for row in rows[1:]:
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = value
            for paragraph in cells[col_idx].paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if col_idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.add_run(f"Table {table_index}. {caption}")
    run.italic = True
    run.font.size = Pt(9)


def add_figure(doc: Document, image_path: Path, caption: str, figure_index: int, width: float = 6.2) -> None:
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.add_run(f"Figure {figure_index}. {caption}")
    run.italic = True
    run.font.size = Pt(9)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for style_name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_paragraph_text(doc: Document, text: str, label_refs: dict[str, str]) -> None:
    cleaned = strip_tex(replace_cross_references(text, label_refs))
    if not cleaned:
        return
    para = doc.add_paragraph(cleaned)
    para.paragraph_format.space_after = Pt(6)


def build_docx(tex_path: Path, bbl_path: Path, output_path: Path) -> None:
    tex = read_text(tex_path)
    label_refs = collect_label_refs(tex, tex_path)
    key_to_num, references = parse_bib_order(bbl_path)
    tex = replace_citations(tex, key_to_num)
    tex = re.sub(r"\\documentclass.*?\\begin\{document\}", "", tex, flags=re.S)
    tex = re.sub(r"\\bibliographystyle\{[^{}]+\}\s*\\bibliography\{[^{}]+\}\s*\\end\{document\}", "", tex, flags=re.S)

    title_match = re.search(r"\\title\{([^{}]+)\}", read_text(tex_path))
    author_match = re.search(r"\\author\{([^{}]+)\}", read_text(tex_path))
    title = title_match.group(1) if title_match else "Submission"
    author = author_match.group(1) if author_match else "Anonymous Authors"

    doc = Document()
    configure_doc(doc)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(16)

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(author)
    author_run.font.name = "Times New Roman"
    author_run.font.size = Pt(12)

    tex = re.sub(r"\\maketitle", "", tex)
    tex = re.sub(r"\\title\{[^{}]+\}|\\author\{[^{}]+\}|\\date\{[^{}]*\}", "", tex)
    tex = tex.replace("\\begin{abstract}", "\n\\section*{Abstract}\n")
    tex = tex.replace("\\end{abstract}", "\n")

    table_index = 1
    figure_index = 1
    lines = tex.splitlines()
    buffer: list[str] = []
    in_figure = False
    figure_block: list[str] = []

    def flush() -> None:
        nonlocal buffer
        text = " ".join(part.strip() for part in buffer if part.strip())
        buffer = []
        add_paragraph_text(doc, text, label_refs)

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush()
            continue
        if stripped.startswith("\\begin{figure}"):
            flush()
            in_figure = True
            figure_block = [stripped]
            continue
        if in_figure:
            figure_block.append(stripped)
            if stripped.startswith("\\end{figure}"):
                block = "\n".join(figure_block)
                img_match = re.search(r"\\includegraphics(?:\[[^\]]+\])?\{([^{}]+)\}", block)
                cap_match = re.search(r"\\caption\{(.*?)\}", block, re.S)
                if img_match and cap_match:
                    image_path = (tex_path.parent / img_match.group(1)).resolve()
                    add_figure(doc, image_path, strip_tex(cap_match.group(1)), figure_index)
                    figure_index += 1
                in_figure = False
            continue
        input_match = re.match(r"\\input\{([^{}]+)\}", stripped)
        if input_match:
            flush()
            input_key = input_match.group(1)
            if input_key in FIGURE_INPUTS:
                fig = FIGURE_INPUTS[input_key]
                add_figure(doc, fig["image"], fig["caption"], figure_index)
                figure_index += 1
            elif input_key in TABLE_INPUTS:
                add_table(doc, TABLE_INPUTS[input_key], table_index)
                table_index += 1
            continue
        sec_match = re.match(r"\\section\*?\{([^{}]+)\}", stripped)
        if sec_match:
            flush()
            doc.add_heading(strip_tex(sec_match.group(1)), level=1)
            continue
        para_match = re.match(r"\\paragraph\{([^{}]+)\}", stripped)
        if para_match:
            flush()
            doc.add_heading(strip_tex(para_match.group(1)), level=3)
            rest = stripped[para_match.end() :].strip()
            if rest:
                buffer.append(rest)
            continue
        if stripped.startswith("\\") and not stripped.startswith("\\text"):
            continue
        buffer.append(stripped)

    flush()

    if references:
        doc.add_heading("References", level=1)
        for idx, entry in enumerate(references, start=1):
            para = doc.add_paragraph(f"[{idx}] {entry}")
            para.paragraph_format.space_after = Pt(3)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.author = "Anonymous Authors"
    doc.core_properties.last_modified_by = "Anonymous Authors"
    doc.core_properties.title = title
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--tex", required=True)
    parser.add_argument("--bbl", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    build_docx(Path(args.tex), Path(args.bbl), Path(args.output))


if __name__ == "__main__":
    main()
